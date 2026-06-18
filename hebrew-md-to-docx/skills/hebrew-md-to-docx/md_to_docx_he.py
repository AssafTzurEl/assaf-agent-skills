#!/usr/bin/env python3
"""
Hebrew Markdown -> Word (.docx) converter, RTL-aware.

Usage:
    python md_to_docx_he.py INPUT.md [-o OUTPUT.docx] [--overrides overrides.json]
                            [--code-font Consolas]

Design principles
-----------------
* Use Word's default template for everything (Normal, Heading 1-N, List Number,
  List Bullet, Quote, Table Grid). The ONLY added style is a "Code" character
  style: identical to Normal but with a monospace font.
* Direction is set per paragraph with w:bidi (RTL paragraphs get NO explicit jc,
  so Word's default right alignment for RTL applies) and per run with w:rtl.
* Runs are split by script and tagged RTL/LTR. Neutral characters (spaces,
  parens, quotes, dashes, punctuation) are tagged LTR only when they sit between
  two LTR characters (so "format string" stays together); otherwise they join
  the RTL/Hebrew flow (so the parens/quotes/dashes around English stay on the
  Hebrew side). Digits are treated as LTR so numbers like 1920 never reverse.
* No Unicode control characters are inserted; nothing visible is added.

Paragraph direction is chosen by the *majority of strong directional characters*
(Hebrew vs Latin). Near-even paragraphs are reported as "ambiguous" and can be
resolved with an overrides file keyed by paragraph index:
    { "118": "rtl", "169": "ltr" }
"""

import argparse
import json
import os
import re
import sys

import mistune
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn

# ---------------------------------------------------------------------------
# Direction detection
# ---------------------------------------------------------------------------

HEB_RE = re.compile(r"[֐-׿יִ-ﭏ]")
LAT_RE = re.compile(r"[A-Za-zÀ-ɏ]")

AMBIGUOUS_BAND = 0.20  # |heb-lat|/total below this -> flag for human review


def counts(text):
    return len(HEB_RE.findall(text)), len(LAT_RE.findall(text))


def detect_direction(text, base="rtl"):
    """Return (direction, ambiguous) for a block of text."""
    h, l = counts(text)
    if h == 0 and l == 0:
        return base, False
    if l == 0:
        return "rtl", False
    if h == 0:
        return "ltr", False
    total = h + l
    direction = "rtl" if h >= l else "ltr"
    ambiguous = abs(h - l) / total < AMBIGUOUS_BAND
    return direction, ambiguous


def _char_dir(ch):
    if HEB_RE.match(ch):
        return "rtl"
    if LAT_RE.match(ch) or ch.isdigit():
        return "ltr"   # digits are LTR so numbers like 1920 never reverse
    return "neutral"


def split_runs(text):
    """Split text into (segment, 'rtl'|'ltr') runs, applying the Unicode bidi
    neutral rule under an RTL base: a neutral char (space, paren, quote, dash,
    punctuation) is LTR only when it lies between two LTR characters; otherwise
    it belongs to the RTL (Hebrew) flow. This keeps English phrases like
    "format string" together while parentheses/quotes/dashes around them stay
    on the Hebrew side.
    """
    if not text:
        return []
    atoms = []  # [segment, dir]
    cur, cur_cls = [], None
    for ch in text:
        cls = _char_dir(ch)
        if cur_cls is None or cls == cur_cls:
            cur.append(ch)
            cur_cls = cls
        else:
            atoms.append(["".join(cur), cur_cls])
            cur, cur_cls = [ch], cls
    atoms.append(["".join(cur), cur_cls])

    n = len(atoms)
    for i, atom in enumerate(atoms):
        if atom[1] != "neutral":
            continue
        prev_dir = next((atoms[j][1] for j in range(i - 1, -1, -1)
                         if atoms[j][1] != "neutral"), None)
        next_dir = next((atoms[j][1] for j in range(i + 1, n)
                         if atoms[j][1] != "neutral"), None)
        atom[1] = "ltr" if (prev_dir == "ltr" and next_dir == "ltr") else "rtl"

    merged = []
    for seg, d in atoms:
        if merged and merged[-1][1] == d:
            merged[-1][0] += seg
        else:
            merged.append([seg, d])
    return [(s, d) for s, d in merged]


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

_PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
]


def _insert_pPr_child(pPr, element):
    tag = element.tag.split("}")[-1]
    idx = _PPR_ORDER.index(tag)
    later = set(_PPR_ORDER[idx + 1:])
    for child in pPr:
        if child.tag.split("}")[-1] in later:
            child.addprevious(element)
            return
    pPr.append(element)


def set_paragraph_direction(paragraph, rtl):
    """Set paragraph base direction via w:bidi only.

    No w:jc is written: an RTL paragraph then defaults to right alignment and an
    LTR paragraph to left alignment, matching the template's defaults.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for b in pPr.findall(qn("w:bidi")):
        pPr.remove(b)
    bidi = OxmlElement("w:bidi")
    if not rtl:
        bidi.set(qn("w:val"), "0")
    _insert_pPr_child(pPr, bidi)


_RPR_ORDER = ["rStyle", "rFonts", "b", "bCs", "i", "iCs", "color", "rtl"]


def _ordered_rPr(children):
    rPr = OxmlElement("w:rPr")
    for tag, el in sorted(children, key=lambda c: _RPR_ORDER.index(c[0])):
        rPr.append(el)
    return rPr


def add_run(paragraph, text, *, bold=False, italic=False, rtl=False, style=None):
    """Add a run. A w:rtl element is written only for RTL (Hebrew) runs; LTR
    runs omit it, exactly as Word itself does (Latin/digits resolve LTR under
    the RTL paragraph base on their own)."""
    r = OxmlElement("w:r")
    children = []
    if style:
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), style)
        children.append(("rStyle", rStyle))
    if bold:
        children.append(("b", OxmlElement("w:b")))
        children.append(("bCs", OxmlElement("w:bCs")))
    if italic:
        children.append(("i", OxmlElement("w:i")))
        children.append(("iCs", OxmlElement("w:iCs")))
    if rtl:
        children.append(("rtl", OxmlElement("w:rtl")))
    r.append(_ordered_rPr(children))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    paragraph._p.append(r)
    return r


def add_break(paragraph):
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:br"))
    paragraph._p.append(r)


# ---------------------------------------------------------------------------
# Inline rendering
# ---------------------------------------------------------------------------

def inline_text(tokens):
    out = []
    for tok in tokens or []:
        t = tok.get("type")
        if t == "text":
            out.append(tok.get("raw", ""))
        elif t == "codespan":
            out.append(tok.get("raw", ""))
        elif t in ("strong", "emphasis", "link", "strikethrough"):
            out.append(inline_text(tok.get("children")))
        elif t in ("softbreak", "linebreak"):
            out.append(" ")
    return "".join(out)


def render_inline(paragraph, tokens, bold=False, italic=False):
    for tok in tokens or []:
        t = tok.get("type")
        if t == "text":
            for seg, d in split_runs(tok.get("raw", "")):
                add_run(paragraph, seg, bold=bold, italic=italic,
                        rtl=(d == "rtl"))
        elif t == "strong":
            render_inline(paragraph, tok.get("children"), bold=True, italic=italic)
        elif t == "emphasis":
            render_inline(paragraph, tok.get("children"), bold=bold, italic=True)
        elif t == "strikethrough":
            render_inline(paragraph, tok.get("children"), bold=bold, italic=italic)
        elif t == "codespan":
            # monospace + LTR run; no control chars, so it never reorders
            add_run(paragraph, tok.get("raw", ""), bold=bold, italic=italic,
                    rtl=False, style="Code")
        elif t == "link":
            render_inline(paragraph, tok.get("children"), bold=bold, italic=italic)
        elif t == "linebreak":
            add_break(paragraph)
        elif t == "softbreak":
            add_run(paragraph, " ")


# ---------------------------------------------------------------------------
# Block rendering
# ---------------------------------------------------------------------------

class Converter:
    def __init__(self, doc):
        self.doc = doc
        self.para_index = 0
        self.ambiguous = []

    def direction_for(self, text, overrides):
        idx = self.para_index
        ovr = overrides.get(str(idx))
        direction, ambiguous = detect_direction(text, base="rtl")
        if ovr in ("rtl", "ltr"):
            direction = ovr
        elif ambiguous:
            self.ambiguous.append((idx, direction, text[:60]))
        self.para_index += 1
        return direction

    def render_blocks(self, tokens, para_style=None):
        for tok in tokens or []:
            self.render_block(tok, para_style)

    def render_block(self, tok, para_style=None):
        t = tok.get("type")
        if t == "heading":
            self.render_heading(tok)
        elif t in ("paragraph", "block_text"):
            self.render_paragraph(tok, para_style)
        elif t == "block_code":
            self.render_code_block(tok)
        elif t == "block_quote":
            self.render_quote(tok)
        elif t == "list":
            self.render_list(tok)
        elif t == "table":
            self.render_table(tok)
        # thematic_break / blank_line: rely on default paragraph flow

    def render_heading(self, tok):
        level = tok["attrs"]["level"]
        text = inline_text(tok["children"])
        direction = self.direction_for(text, self.overrides)
        p = self.doc.add_paragraph(style=f"Heading {min(level, 9)}")
        set_paragraph_direction(p, direction == "rtl")
        render_inline(p, tok["children"])

    def render_paragraph(self, tok, para_style=None):
        text = inline_text(tok["children"])
        direction = self.direction_for(text, self.overrides)
        p = self.doc.add_paragraph(style=para_style)
        set_paragraph_direction(p, direction == "rtl")
        render_inline(p, tok["children"])
        return p

    def render_code_block(self, tok):
        raw = tok.get("raw", "").rstrip("\n")
        lines = raw.split("\n")
        p = self.doc.add_paragraph()
        set_paragraph_direction(p, rtl=False)
        for i, line in enumerate(lines):
            add_run(p, line if line else " ", rtl=False, style="Code")
            if i != len(lines) - 1:
                add_break(p)
        self.para_index += 1

    def render_quote(self, tok):
        style = "Quote" if "Quote" in (s.name for s in self.doc.styles) else None
        self.render_blocks(tok.get("children"), para_style=style)

    def _new_num(self, ordered):
        """Allocate a fresh numId (same abstract list) so each list restarts at 1."""
        abs_id = _ABS_ORDERED if ordered else _ABS_BULLET
        nid = self.next_num_id
        self.next_num_id += 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(nid))
        ai = OxmlElement("w:abstractNumId")
        ai.set(qn("w:val"), str(abs_id))
        num.append(ai)
        lo = OxmlElement("w:lvlOverride")
        lo.set(qn("w:ilvl"), "0")
        so = OxmlElement("w:startOverride")
        so.set(qn("w:val"), "1")
        lo.append(so)
        num.append(lo)
        self.numbering_el.append(num)
        return nid

    def render_list(self, tok):
        ordered = tok["attrs"].get("ordered", False)
        num_id = self._new_num(ordered)
        for item in tok.get("children", []):
            first = True
            for child in item.get("children", []):
                if child.get("type") in ("block_text", "paragraph"):
                    p = self.render_paragraph(child, para_style=list_paragraph_style(self.doc))
                    if first:
                        apply_numbering(p, num_id)
                        first = False
                else:
                    self.render_block(child)

    def render_table(self, tok):
        head = None
        body_rows = []
        for child in tok.get("children", []):
            if child["type"] == "table_head":
                head = child["children"]
            elif child["type"] == "table_body":
                body_rows = child["children"]
        ncols = len(head) if head else (len(body_rows[0]["children"]) if body_rows else 0)
        if ncols == 0:
            return
        table = self.doc.add_table(rows=0, cols=ncols)
        set_table_borders(table)
        table._tbl.tblPr.append(OxmlElement("w:bidiVisual"))
        if head:
            row = table.add_row()
            for cell, cellfmt in zip(row.cells, head):
                self._fill_cell(cell, cellfmt, header=True)
        for r in body_rows:
            row = table.add_row()
            for cell, cellfmt in zip(row.cells, r["children"]):
                self._fill_cell(cell, cellfmt, header=False)

    def _fill_cell(self, cell, cellfmt, header):
        p = cell.paragraphs[0]
        for r in list(p._p.findall(qn("w:r"))):
            p._p.remove(r)
        text = inline_text(cellfmt.get("children"))
        direction, _ = detect_direction(text, base="rtl")
        set_paragraph_direction(p, direction == "rtl")
        render_inline(p, cellfmt.get("children"), bold=header)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def ensure_code_style(doc, code_font):
    if "Code" in (s.name for s in doc.styles):
        return
    st = doc.styles.add_style("Code", WD_STYLE_TYPE.CHARACTER)
    st.base_style = doc.styles["Default Paragraph Font"] if \
        "Default Paragraph Font" in (s.name for s in doc.styles) else None
    st.font.name = code_font
    rPr = st.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:cs"), code_font)
    rFonts.set(qn("w:hAnsi"), code_font)
    rFonts.set(qn("w:ascii"), code_font)


def set_table_borders(table, color="auto", sz="4"):
    """Single borders on all sides + inside grid, without depending on the
    'Table Grid' style (which may be absent from a custom template)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def list_paragraph_style(doc):
    return "List Paragraph" if "List Paragraph" in (s.name for s in doc.styles) else None


NUMID_ORDERED = 1101
NUMID_BULLET = 1102
_ABS_ORDERED = 1101
_ABS_BULLET = 1102

_NUMBERING_XML = f"""<w:numbering xmlns:w="{nsmap['w']}">
  <w:abstractNum w:abstractNumId="{_ABS_ORDERED}">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:start="720" w:hanging="360"/></w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="{_ABS_BULLET}">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="•"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:start="720" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="{NUMID_ORDERED}"><w:abstractNumId w:val="{_ABS_ORDERED}"/></w:num>
  <w:num w:numId="{NUMID_BULLET}"><w:abstractNumId w:val="{_ABS_BULLET}"/></w:num>
</w:numbering>"""


def ensure_numbering(doc):
    """Attach a numbering part with one decimal and one bullet definition, so
    lists work regardless of whether the base template defines list styles."""
    from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.oxml import parse_xml
    from docx.parts.numbering import NumberingPart
    try:
        doc.part.numbering_part  # already present
        return
    except NotImplementedError:
        pass
    part = NumberingPart(PackURI("/word/numbering.xml"), CT.WML_NUMBERING,
                         parse_xml(_NUMBERING_XML), doc.part.package)
    doc.part.relate_to(part, RT.NUMBERING)


def apply_numbering(paragraph, num_id, ilvl=0):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(qn("w:val"), str(ilvl)); numPr.append(il)
    ni = OxmlElement("w:numId"); ni.set(qn("w:val"), str(num_id)); numPr.append(ni)
    pPr.insert(0, numPr)


def strip_leading_empty(doc):
    """Drop blank paragraphs a template may start with, so content begins clean."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            if child.find(qn("w:r")) is None and not (child.text or "").strip():
                body.remove(child)
            else:
                break
        else:
            break


def convert(md_path, out_path, overrides, code_font, template=None):
    with open(md_path, encoding="utf-8") as f:
        text = f.read()
    md = mistune.create_markdown(
        renderer=None,
        plugins=["table", "strikethrough", "url"],
    )
    tokens = md(text)
    # Base the document on a template if given (or a bundled template.docx next
    # to this script); otherwise use python-docx's stock default. The converter
    # is template-agnostic: it supplies its own numbering and table borders and
    # only relies on Normal / Heading N / Quote, which any Word template has.
    if template is None:
        bundled = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "template.docx")
        if os.path.exists(bundled):
            template = bundled
    doc = Document(template) if template else Document()
    ensure_code_style(doc, code_font)
    ensure_numbering(doc)
    conv = Converter(doc)
    conv.overrides = overrides
    conv.numbering_el = doc.part.numbering_part.element
    conv.next_num_id = 2000
    conv.render_blocks(tokens)
    strip_leading_empty(doc)
    doc.save(out_path)
    return conv


def main():
    ap = argparse.ArgumentParser(description="Hebrew Markdown -> DOCX (RTL-aware)")
    ap.add_argument("input")
    ap.add_argument("-o", "--output")
    ap.add_argument("--overrides", help="JSON map of paragraph index -> rtl/ltr")
    ap.add_argument("--code-font", default="Consolas")
    ap.add_argument("--template", help="base .docx to inherit styles from "
                    "(e.g. a blank doc saved from your Word default template)")
    args = ap.parse_args()

    out = args.output or re.sub(r"\.md$", "", args.input) + ".docx"
    overrides = {}
    if args.overrides:
        with open(args.overrides, encoding="utf-8") as f:
            overrides = json.load(f)

    conv = convert(args.input, out, overrides, args.code_font, args.template)
    print(f"Wrote {out}")
    if conv.ambiguous:
        print(f"\n{len(conv.ambiguous)} ambiguous paragraph(s) "
              f"(near-even Hebrew/Latin counts). Resolve via --overrides "
              f"(index -> rtl/ltr):", file=sys.stderr)
        for idx, d, snippet in conv.ambiguous:
            print(f"  [{idx}] used={d}: {snippet!r}", file=sys.stderr)


if __name__ == "__main__":
    main()
